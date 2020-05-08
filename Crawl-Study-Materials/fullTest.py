from selenium import webdriver
import re
import docx
from docx import Document
import time


def writeWord(list):
    document = Document()
    for x in range(len(list)):
        document.add_paragraph(list[x]["testLine"])
        document.add_paragraph(list[x]["answer"])
        document.add_paragraph()

    document.save('full.docx')



def getData(brower):
    allInfos = []

    fullpage = brower.find_element_by_xpath('//*[@id="sectionRender"]/div[3]/div[2]/a[1]')
    fullpage.click()

    time.sleep(1)
    for i in range(25):
        metaInfos = {}

        time.sleep(1)
        full_1 = brower.find_element_by_xpath(
            '//*[@id="questionRender"]/div[2]/div[2]/div[3]/div[2]/div[2]/div[1]/input'
        )
        full_1.send_keys("1")
        time.sleep(1)
        outAnswer = brower.find_element_by_xpath(
            '//*[@id="questionRender"]/div[2]/div[2]/div[4]/a[2]'
        )
        outAnswer.click()


        testLineTitle = brower.find_element_by_xpath(
            '//*[ @ id = "questionRender"] / div[2] / div[2] / div[3] / div[1]/div[1]'
        )
        testLineMes = brower.find_element_by_xpath(
            '//*[ @ id = "questionRender"] / div[2] / div[2] / div[3] / div[1]/div[2]'
        )
        testLine = str(testLineTitle.text) + str(testLineMes.text)
        metaInfos['testLine'] = testLine
       # print(testLine)

        '''
        for n in range(1, 5):
            choiceTitle = brower.find_element_by_xpath(
                '//*[@id="questionRender"]/div[2]/div[2]/div[3]/div[2]/label[' + str(n) + ']/div[1]'
            )
            choiceMes = brower.find_element_by_xpath(
                '//*[@id="questionRender"]/div[2]/div[2]/div[3]/div[2]/label[' + str(n) + ']/div[2]'
            )
            choice = str(choiceTitle.text) + str(choiceMes.text)
            if n==1:
                metaInfos['A'] = choice
            elif n==2:
                metaInfos['B'] = choice
            elif n==3:
                metaInfos['C'] = choice
            else:
                metaInfos['D'] = choice
                #print(metaInfos['D'])
        '''

        answerTitle = brower.find_element_by_xpath(
            '//*[@id="questionRender"]/div[3]/div[2]/div[1]/div[2]/div/div[1]'
        )
        answerMes = brower.find_element_by_xpath(
            '//*[@id="questionRender"]/div[3]/div[2]/div[1]/div[2]/div/div[2]'
        )
        answer = str(answerTitle.text) + str(answerMes.text)
        print(answer)
        metaInfos['answer'] = answer
        allInfos.append(metaInfos)



        # print()
        if i == 24:
            break
        nextPage = brower.find_element_by_xpath('//*[@id="sectionRender"]/div[3]/div[2]/a[' + str(i+2) + ']')
        nextPage.click()

    #print(allInfos)

    return allInfos

brower = webdriver.Chrome()
brower.get("http://saishi.cnki.net/Exercise/Practice/ks0af38326-10ce-48a5-a2cb-84204baf254c/1")
writeWord(getData(brower))
